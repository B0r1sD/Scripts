#!/usr/bin/python3
#Boris Depoortere, 19/01/2023, Scripting exam pt. 1

#importing modules

import MySQLdb as my
from docx import Document
from docx.shared import Cm
import matplotlib.pyplot as plt
import numpy as np


###############################################################
# A Input
###############################################################
print("This script will look for transcript data using Ensembl ENGS accession numbers")
print("="*62)
accessions = []

# if user enters 0 --> the input ask will stop and the script will continue
while True:
    acc = input("Enter accession no. or type 0 to stop: ")
    if acc != "0":
        accessions.append(acc)
        continue
    else:
        break
print("\nYour list:\n")
print(accessions)


###############################################################
# B data handling 
###############################################################

# Opening a blank document based on default template
document = Document()

# connect to UCSC
try:
    conn = my.connect(host="genome-euro-mysql.soe.ucsc.edu",
        user="genomep",
        passwd="password",
        db="hg19")
    cursor = conn.cursor()
except my.OperationalError:
    print("\nConnection to UCSC database failed. Check your internet connection and try again.\n")

# example accessions: ENSG00000223972, ENSG00000137757, ENSG00000169862 and ENSG00000144283
print("Processing Ensembl gene list:")
print("="*62)

#plot data list opener
TranscriptList = []
totalTranscripts = []

for accession in accessions:
    query = "SELECT * FROM wgEncodeGencodeAttrsV36lift37 WHERE geneId LIKE "
    q = """{}'{}%'""".format(query,accession)

    #visual output
    print("Getting data for {} using query:".format(accession))
    print("\t{}".format(q))
    transcript = cursor.execute(q)
    print("Found {} transcripts\n".format(transcript))
    
    #generate plot data list
    totalTranscripts.append(transcript)
    
    
    
    # commit
    conn.commit()
    
    # get all results of the SNP query
    result = cursor.fetchall()
    
    #print(result)
    
    rowNr = 0
    
    #iterating over result
    for row in result:
        if rowNr == 0:
            ############# creating docx header ############
            
            
            #assigning names to asked information for readability (FAIR)
            gsymbol = row[1]
            transcriptID = row[4]
            ttype = row[6]
            level = row[11]
            
            ####### building Word file .doxc #########
            
            document.add_heading('{}'.format(accession),1)

            table = document.add_table(rows=1, cols=4)
            table.style = 'LightShading-Accent1'
            
            # Header row
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'GENE SYMBOL'
            hdr_cells[1].text = 'TRANSCRIPT ID'
            hdr_cells[2].text = 'TRANSCRIPT TYPE'
            hdr_cells[3].text = 'LEVEL'
            
            #ending header creation
            rowNr += 1
        else:
            # Add data rows
            #print(row)
            
            row_cells = table.add_row().cells
                    
            row_cells[0].text = gsymbol
            row_cells[1].text = transcriptID
            row_cells[2].text = ttype
            row_cells[3].text = str(level)
            
        
        #creating transcript type list
        
        TranscriptList.append(ttype)
    RowNr = 0
    document.add_page_break()

#################barplot creation##################

#unique values --> creating set

#print(TranscriptList)
YAxis = set(TranscriptList)
#print(YAxis)

UniqueTranscripts = list(YAxis)
UniqueTranscripts.sort()

print(UniqueTranscripts)
#hardcoded bcs sometimes not same amount of x and y ... above works for 1, 2 examples but for 4 there is an inbalance
#UniqueTranscripts = ['nonsense_mediated_decay', 'processed_transcript', 'protein_coding','filler','filler'']

print(totalTranscripts)

plt.figure(figsize=(10,6))
plt.xlabel("Total number of transcripts",fontsize=12)
plt.title("Transcripts per type")
plt.barh(UniqueTranscripts, totalTranscripts,align='center')


#tick labels
from matplotlib.ticker import StrMethodFormatter
plt.gca().xaxis.set_major_formatter(StrMethodFormatter('{x:,.0f}'))
plt.savefig('./answer1.png', bbox_inches='tight')
plt.show()


#adding picture at end of page

document.add_picture('answer1.png', width=Cm(12))

        
#saving files and closing connection
conn.close()
document.save('answer1.docx')

#print(totalTranscripts)

# example accessions: ENSG00000223972, ENSG00000137757, ENSG00000169862 and ENSG00000144283