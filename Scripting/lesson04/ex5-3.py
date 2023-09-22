#!/usr/bin/python

import os
import re
import fileinput
from docx import Document
from docx.shared import Inches,Cm

document = Document()
document.add_heading('FastQC report (via Python)')
#document.add_picture('')

location = "/home/guest/BITsemA/Scripting/lesson04/fastqc"
os.chdir(location)

print("current working directory: {}".format(os.getcwd()))


'''
for x in list_dir:
    for y in x:
        print(y)'''

i = 0

for root, dirs, files in os.walk(location):
   #print(root,dirs,files)
    
   #print(files)
    for x in dirs:
        if re.search("[^Icons,Images]", x) != None:
            listname = x.split("_")
            
            # list_dir = os.listdir(location)
    #print(list_dir[i])
    i += 1

    '''if re.search("[^.html]", list_dir[i]) != None:
        print(list_dir[i])'''
        #print(files)
        #basename = os.path.basename(files)
        #print(basename)

    for name in files:
        
        if re.search("fastqc_data.txt", name) != None:
            textfile = fileinput.input(root + "/" + name)
            '''for x in textfile:
                print(x)'''
            for x in textfile:
                if re.search("Total Sequences", x) != None:
                    #print(x)
                    document.add_heading("Sample " + listname[0],level=1)
                    p = document.add_heading('summary',level=2)
                    document.add_paragraph(x)
        if re.search("summary.txt", name) != None:
            sumfile = fileinput.input(root + "/" + name)
            for x in sumfile:
                if re.search("Basic Statistics", x) != None:
                    list1 = x.split("\t")
                    #print(list1)
                    '''listname = list1[2].split(".")
                    document.add_heading(listname[0],level=1)'''
                    document.add_paragraph(list1[0] + "\t" + list1[1])
                if re.search("Per base sequence quality", x) != None:
                    #print(x)
                    list2 = x.split("\t")
                    document.add_paragraph(list2[0] + "\t" + list2[1])
        if re.search("per_base_quality.png", name) != None:
           #print(os.path.join(root, name))
           document.add_picture(os.path.join(root, name),width=Cm(15))
           document.add_page_break()
    

      #print(os.path.join(root, name))
os.chdir('/home/guest/BITsemA/Scripting/lesson04')
document.save('FastQC_results_parsed.docx')

'''for name in dirs:
      print(os.path.join(root, name))'''