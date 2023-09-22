#!/usr/bin/python3

import os
import sqlite3
import MySQLdb as mydb
import regex as re
import matplotlib.pyplot as plt
import numpy as np

#connect to UCSC 1
print("\nREADING DATA FROM UCSC DATABASE")
try: 
    db = mydb.connect(host="genome-mysql.soe.ucsc.edu",
   user="genomep",
   passwd="password",
   db="hg19")

except sqlite3.OperationalError:
    print('Nope UCSC')

c = db.cursor()

from docx import Document
# Opening a blank document based on default template
document = Document()

yax = []

with open('answer.csv', 'r') as f:
    print("Reading regions from file:")
    print("*"*30)
    for line in f:
        #print(line)
        indiv = line.strip().split(",")
        startstop = indiv[0].split(":")
        
        startstop1 = startstop[1].split("-")
        print("Chrom = {}".format(startstop[0]))
        print("Query:")
        query = "SELECT * FROM targetScanS WHERE chrom = '{}' AND chromStart > '{}' AND chromEnd < '{}'".format(startstop[0],startstop1[0],startstop1[1])
        print(query)
        c.execute(query)
        results = c.fetchall()
        print("results: {}".format(len(results)))
        yax.append(len(results))
        len(results)
        #print(results)
        db.commit()

        #******************building docx*****************
        document.add_heading('TS miRNA sites for {}'.format(indiv[0]), 2)

        table = document.add_table(rows=1, cols=5)
        table.style = 'LightShading-Accent1'
        
        # Header row
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'name gene'
        hdr_cells[1].text = 'name miRNA'
        hdr_cells[2].text = 'position'
        hdr_cells[3].text = 'score'
        hdr_cells[4].text = 'strand'
        # Add data rows
        
        
        for indiv1 in results:

            row_cells = table.add_row().cells
            
            sole_gene = re.findall('\w*:',indiv1[4])
            
            row_cells[0].text = sole_gene
            
            sole_mirna = re.findall(':\w*',indiv1[4])
            row_cells[1].text = sole_mirna
            
            row_cells[2].text = str(line)
            row_cells[-1].text = indiv1[-1]
            #row_cells[3].text = indiv1[2]
            
        document.add_page_break()
    # plot
    print(yax)
    plt.xlabel("Total number of sites",fontsize=14)
    plt.title("plot 1")
    plt.legend()
    plt.barh(line, yax,align='center')
        
# end open file

db.close()
document.save('example.docx')



############### matplotlib #################


'''
plt.figure(figsize=(10,6))

plt.xlabel("Total number of sites",fontsize=14)
plt.grid(True)
plt.bar([0,5,10,15,20],len(results), align='center',tick_label=line)
'''



plt.show()